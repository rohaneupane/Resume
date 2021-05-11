from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor

FONT = 'Arial'
document = Document()
sections = document.sections

for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def format_page(before_space, after_space, spacing_line, page_obj):
    page_obj.paragraph_format.space_before = Inches(before_space)
    page_obj.paragraph_format.space_after = Inches(after_space)
    page_obj.paragraph_format.line_spacing = spacing_line


def main_heading(text):
    page_obj = document.add_paragraph()
    page_obj.style = document.styles['Heading 1']
    document.add_picture("Horizontal line.png", width=Inches(7.8), height=Inches(0.25))
    run_page_obj = page_obj.add_run(text)
    run_page_obj.font.color.rgb = RGBColor(0, 0, 0)


def edu(text_1, text_2, text_3, text_4, text_5):
    line_1 = document.add_paragraph()
    line_1.add_run(f"{text_1}, ").bold = True
    line_1.add_run(text_2)
    line_1.add_run(text_5).italic = True
    line_1 = document.add_paragraph()
    line_1.add_run(text_3)
    format_page(0, 0, 1, line_1)
    line_1 = document.add_paragraph(text_4, style='List Bullet 2')


def head_section_2(heading):
    body_2 = document.add_paragraph()
    run_body_2 = body_2.add_run(heading)
    run_body_2.bold = True
    run_body_2.font.size = Pt(12)
    return body_2


def points():
    text = input("Enter your point: ")
    return document.add_paragraph(text, style='List Bullet 2')


def experience_sub(obj, text, year):
    obj.add_run(f" - {year}").italic = True
    format_page(0, 0, 1, obj)
    obj = document.add_paragraph()
    run_obj = obj.add_run(text).italic = True


def all_skills():
    for skill in skill_sec:
        skill_set = document.add_paragraph(style='List Bullet 2')
        skill_set.add_run(f"{skill}: ").bold = True
        skill_set.add_run(input("Enter all your skills separated by commas: "))


#Top-Heading for Name
heading = document.add_paragraph()
heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
format_page(0, 0, 1, heading)
run_heading = heading.add_run(input("What is your name? "))
run_heading.bold = True
run_heading.font.size = Pt(22)
run_heading.font.name = FONT

#Sub-Heading for details
sub_heading = document.add_paragraph()
sub_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_sub_heading = sub_heading.add_run(f"{input('Enter your email: ')} | +91 {input('Enter your number: ')} | {input('Your City: ')}"
                                      f"-{input('Pincode: ')}, {input('Your State: ')}")
run_sub_heading.font.size = Pt(11)
run_sub_heading.font.name = FONT

##################THIS IS THE MAIN CODE##############################################
#Section 1
section_1 = main_heading("EDUCATION")
for _ in range(0,int(input("How many education headings do you want? "))):
    edu(input("University Name: "), input("Institute Name: "), input("What course: "),
        input("Marks obtained, Percentage or GPA: "), f" - {input('Year of Completion: ')}")


#Section 2
section_2 = main_heading("EXPERIENCE")
for _ in range(0, int(input("How many experience headings do you want? "))):
    heading = head_section_2(input("Your Heading: "))
    experience_sub(heading, input("Your Subheading: "), input("Which year did you do it? "))
    for _ in range(0, int(input("How many points do you want? "))):
        points()

#Section 3
section_3 = main_heading("SKILLS")
skill_sec = []
for _ in range(0, int(input("How many skill sections do you want? "))):
    skill_sec.append(input("Enter the skill section: "))
all_skills()
##################THIS IS THE MAIN CODE##############################################


end = document.add_paragraph("This Resume was made using python-docx package")
end.paragraph_format.space_before = Inches(2)
document.save("Resume.docx")


